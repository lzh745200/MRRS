"""报表导出服务 — Word/PDF 公文报告生成。

数据约定：各 ``generate_*_report_data`` 方法返回统一结构::

    {"year": int, "title": str, "sections": [Section]}

其中 ``Section = {"title": str, "paragraphs": [str], "table": {"headers": [...], "rows": [[...]]} | None}``，
``export_word`` / ``export_pdf`` 对该结构做通用渲染，空数据时输出"暂无数据"说明而非空文件。
"""

import logging
from datetime import datetime
from io import BytesIO

logger = logging.getLogger(__name__)

# 与 export.py 的 _REPORT_TYPE_MAP 保持一致（service 层独立维护，避免反向依赖 api 层）
_REPORT_TITLES = {
    "summary": "年度帮扶工作总结",
    "fund_detail": "帮扶资金拨付明细表",
    "project_progress": "帮扶项目进度统计表",
    "school_statistics": "帮扶学校统计表",
    "village_summary": "帮扶村年度汇总表",
    "annual_summary": "年度综合总结报告",
}

# 项目状态中文映射（对应 models/project.py ProjectStatus）
_PROJECT_STATUS_CN = {
    "draft": "草稿",
    "pending": "待审批",
    "approved": "已批准",
    "in_progress": "进行中",
    "completed": "已完成",
    "cancelled": "已取消",
}


def _fmt_amount(value) -> str:
    """金额展示格式化（千元分隔，保留两位小数）。"""
    try:
        return f"{float(value):,.2f}"
    except (TypeError, ValueError):
        return "0.00"


class ReportExportService:
    """报表导出服务（单例）"""

    # ────────────────────── 数据构建 ──────────────────────

    @staticmethod
    def _resolve_year(year) -> int:
        return int(year) if year else datetime.now().year

    def generate_summary_report_data(self, db, year: int) -> dict:
        """生成年度帮扶工作总结数据（村/校/项目/经费四块合并）"""
        year = self._resolve_year(year)
        village = self.generate_village_summary_report_data(db, year)
        school = self.generate_school_statistics_report_data(db, year)
        project = self.generate_project_progress_report_data(db, year)
        fund = self.generate_fund_detail_report_data(db, year)

        sections = [
            {
                "title": "一、帮扶村概况",
                "paragraphs": [f"截至报告期末，在册帮扶村共 {village['total_villages']} 个。"],
                "table": None,
            },
            {
                "title": "二、帮扶学校概况",
                "paragraphs": [
                    f"在册帮扶学校 {school['total_schools']} 所，"
                    f"在校学生 {school['total_students']} 人，教职工 {school['total_teachers']} 人。"
                ],
                "table": None,
            },
            {
                "title": "三、帮扶项目进展",
                "paragraphs": [],
                "table": project["sections"][0]["table"] if project["sections"] else None,
            },
            {
                "title": "四、帮扶资金执行",
                "paragraphs": [],
                "table": fund["sections"][0]["table"] if fund["sections"] else None,
            },
        ]
        return {"year": year, "title": _REPORT_TITLES["summary"], "sections": sections}

    def generate_fund_detail_report_data(self, db, year: int) -> dict:
        """生成经费明细报表数据（按经费类型分组聚合）"""
        from sqlalchemy import func

        from app.models.fund import Fund

        year = self._resolve_year(year)
        try:
            year_col = func.extract("year", func.coalesce(Fund.date, Fund.created_at))
            type_col = func.coalesce(Fund.fund_type, "未分类")
            rows = (
                db.query(
                    type_col.label("ftype"),
                    func.count(Fund.id),
                    func.coalesce(func.sum(Fund.amount), 0),
                    func.coalesce(func.sum(Fund.approved_amount), 0),
                    func.coalesce(func.sum(Fund.allocated_amount), 0),
                    func.coalesce(func.sum(Fund.used_amount), 0),
                )
                .filter(Fund.is_active == True, year_col == year)  # noqa: E712
                .group_by(type_col)
                .order_by(type_col)
                .all()
            )

            table_rows = [
                [r[0], str(r[1]), _fmt_amount(r[2]), _fmt_amount(r[3]), _fmt_amount(r[4]), _fmt_amount(r[5])]
                for r in rows
            ]
            totals = [
                "合计",
                str(sum(r[1] for r in rows)),
                _fmt_amount(sum(r[2] for r in rows)),
                _fmt_amount(sum(r[3] for r in rows)),
                _fmt_amount(sum(r[4] for r in rows)),
                _fmt_amount(sum(r[5] for r in rows)),
            ]
            if rows:
                table_rows.append(totals)

            table = {
                "headers": ["经费类型", "笔数", "申请金额(元)", "批准金额(元)", "拨付金额(元)", "使用金额(元)"],
                "rows": table_rows,
            }
            paragraphs = [] if rows else [f"{year} 年度暂无经费记录。"]
            return {
                "year": year,
                "title": _REPORT_TITLES["fund_detail"],
                "sections": [{"title": "经费执行明细", "paragraphs": paragraphs, "table": table}],
            }
        except Exception:
            logger.warning("经费明细数据查询失败（报表将输出空表）", exc_info=True)
            return {
                "year": year,
                "title": _REPORT_TITLES["fund_detail"],
                "sections": [{"title": "经费执行明细", "paragraphs": [f"{year} 年度经费数据查询失败。"], "table": None}],
            }

    def generate_project_progress_report_data(self, db, year: int) -> dict:
        """生成项目进度报表数据（按状态分组聚合）"""
        from sqlalchemy import func

        from app.models.project import Project

        year = self._resolve_year(year)
        try:
            year_col = func.extract("year", func.coalesce(Project.start_date, Project.created_at))
            rows = (
                db.query(
                    Project.status,
                    func.count(Project.id),
                    func.coalesce(func.avg(Project.progress), 0),
                    func.coalesce(func.sum(Project.budget), 0),
                    func.coalesce(func.sum(Project.actual_cost), 0),
                )
                .filter(year_col == year)
                .group_by(Project.status)
                .order_by(Project.status)
                .all()
            )

            table_rows = [
                [
                    _PROJECT_STATUS_CN.get(r[0], r[0] or "未知"),
                    str(r[1]),
                    f"{float(r[2]):.1f}%",
                    _fmt_amount(r[3]),
                    _fmt_amount(r[4]),
                ]
                for r in rows
            ]
            if rows:
                table_rows.append(
                    [
                        "合计",
                        str(sum(r[1] for r in rows)),
                        "-",
                        _fmt_amount(sum(r[3] for r in rows)),
                        _fmt_amount(sum(r[4] for r in rows)),
                    ]
                )

            table = {
                "headers": ["项目状态", "项目数", "平均进度", "预算金额(元)", "实际花费(元)"],
                "rows": table_rows,
            }
            paragraphs = [] if rows else [f"{year} 年度暂无项目记录。"]
            return {
                "year": year,
                "title": _REPORT_TITLES["project_progress"],
                "sections": [{"title": "项目进展统计", "paragraphs": paragraphs, "table": table}],
            }
        except Exception:
            logger.warning("项目进度数据查询失败（报表将输出空表）", exc_info=True)
            return {
                "year": year,
                "title": _REPORT_TITLES["project_progress"],
                "sections": [{"title": "项目进展统计", "paragraphs": [f"{year} 年度项目数据查询失败。"], "table": None}],
            }

    def generate_school_statistics_report_data(self, db, year: int) -> dict:
        """生成帮扶学校统计报表数据"""
        from sqlalchemy import func

        from app.models.school import School

        year = self._resolve_year(year)
        try:
            schools = db.query(
                func.count(School.id),
                func.coalesce(func.sum(School.student_count), 0),
                func.coalesce(func.sum(School.teacher_count), 0),
            ).filter(School.is_active == True).first()  # noqa: E712

            total_schools = int(schools[0]) if schools else 0
            total_students = int(schools[1]) if schools else 0
            total_teachers = int(schools[2]) if schools else 0
        except Exception:
            # 报表数据查询失败若返回全 0 会生成误导性正式报表，升级为 warning 可观测
            logger.warning("学校统计数据查询失败（报表将输出 0）", exc_info=True)
            total_schools = total_students = total_teachers = 0

        return {
            "year": year,
            "title": _REPORT_TITLES["school_statistics"],
            "total_schools": total_schools,
            "total_students": total_students,
            "total_teachers": total_teachers,
            "sections": [
                {
                    "title": "帮扶学校总体情况",
                    "paragraphs": [
                        f"在册帮扶学校 {total_schools} 所，在校学生 {total_students} 人，"
                        f"教职工 {total_teachers} 人。"
                    ],
                    "table": None,
                }
            ],
        }

    def generate_village_summary_report_data(self, db, year: int) -> dict:
        """生成帮扶村年度汇总报表数据"""
        from sqlalchemy import func

        from app.models.supported_village import SupportedVillage

        year = self._resolve_year(year)
        try:
            total = db.query(func.count(SupportedVillage.id)).filter(
                SupportedVillage.is_active == True  # noqa: E712
            ).scalar() or 0
        except Exception:
            logger.warning("帮扶村汇总数据查询失败（报表将输出 0）", exc_info=True)
            total = 0

        return {
            "year": year,
            "title": _REPORT_TITLES["village_summary"],
            "total_villages": total,
            "sections": [
                {
                    "title": "帮扶村总体情况",
                    "paragraphs": [f"截至报告期末，在册帮扶村共 {total} 个。"],
                    "table": None,
                }
            ],
        }

    def generate_annual_summary_report_data(self, db, year: int) -> dict:
        """生成年度综合总结报告数据（合并村/项目/资金/学校）"""
        year = self._resolve_year(year)
        village_data = self.generate_village_summary_report_data(db, year)
        school_data = self.generate_school_statistics_report_data(db, year)
        project_data = self.generate_project_progress_report_data(db, year)
        fund_data = self.generate_fund_detail_report_data(db, year)

        sections = [
            {"title": "一、帮扶村概况", "paragraphs": village_data["sections"][0]["paragraphs"], "table": None},
            {"title": "二、帮扶学校概况", "paragraphs": school_data["sections"][0]["paragraphs"], "table": None},
            {"title": "三、帮扶项目进展", "paragraphs": [], "table": project_data["sections"][0]["table"]},
            {"title": "四、帮扶资金执行", "paragraphs": [], "table": fund_data["sections"][0]["table"]},
        ]
        return {
            "year": year,
            "title": _REPORT_TITLES["annual_summary"],
            "village_summary": village_data,
            "school_statistics": school_data,
            "project_progress": project_data,
            "fund_detail": fund_data,
            "sections": sections,
        }

    # ────────────────────── 渲染输出 ──────────────────────

    @staticmethod
    def _iter_render_items(data: dict):
        """将统一数据结构展开为渲染项（标题/正文/表格）序列。"""
        for section in data.get("sections") or []:
            yield "heading", section.get("title") or ""
            for para in section.get("paragraphs") or []:
                yield "paragraph", para
            table = section.get("table")
            if table:
                if table.get("rows"):
                    yield "table", table
                else:
                    yield "paragraph", "暂无数据。"

    def export_word(self, report_type: str, data: dict) -> bytes:
        """导出 Word 文档（python-docx）"""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Pt

        title = data.get("title") or _REPORT_TITLES.get(report_type, "帮扶工作报告")

        doc = Document()
        # 中文字体：正文宋体、标题黑体（仅声明字体名，实际字形由打开方系统提供）
        normal = doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(12)
        normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")

        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.name = "SimHei"
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")

        meta = doc.add_paragraph(f"年度：{data.get('year', '-')}    生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

        has_content = False
        for kind, payload in self._iter_render_items(data):
            has_content = True
            if kind == "heading":
                h = doc.add_heading(payload, level=1)
                for run in h.runs:
                    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
            elif kind == "paragraph":
                doc.add_paragraph(payload)
            else:
                headers = payload.get("headers") or []
                rows = payload.get("rows") or []
                table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
                table.style = "Table Grid"
                for j, text in enumerate(headers):
                    table.rows[0].cells[j].text = str(text)
                for i, row in enumerate(rows, start=1):
                    for j, cell in enumerate(row):
                        table.rows[i].cells[j].text = str(cell)

        if not has_content:
            doc.add_paragraph("本年度暂无相关数据。")

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def export_pdf(self, report_type: str, data: dict) -> bytes:
        """导出 PDF 文档（reportlab，内置 CID 中文字体，无需字体文件）"""
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

        if "STSong-Light" not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

        title = data.get("title") or _REPORT_TITLES.get(report_type, "帮扶工作报告")
        title_style = ParagraphStyle(
            "title", fontName="STSong-Light", fontSize=18, leading=26, alignment=1, spaceAfter=6
        )
        meta_style = ParagraphStyle(
            "meta", fontName="STSong-Light", fontSize=10, leading=14, alignment=1, textColor=colors.grey
        )
        heading_style = ParagraphStyle(
            "heading", fontName="STSong-Light", fontSize=14, leading=20, spaceBefore=10, spaceAfter=4
        )
        body_style = ParagraphStyle("body", fontName="STSong-Light", fontSize=11, leading=16)

        story = [
            Paragraph(title, title_style),
            Paragraph(
                f"年度：{data.get('year', '-')}    生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}",
                meta_style,
            ),
            Spacer(1, 6 * mm),
        ]

        has_content = False
        for kind, payload in self._iter_render_items(data):
            has_content = True
            if kind == "heading":
                story.append(Paragraph(payload, heading_style))
            elif kind == "paragraph":
                story.append(Paragraph(payload, body_style))
            else:
                headers = [str(h) for h in (payload.get("headers") or [])]
                rows = [[str(c) for c in row] for row in (payload.get("rows") or [])]
                table = Table([headers, *rows], repeatRows=1)
                table.setStyle(
                    TableStyle(
                        [
                            ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                            ("FONTSIZE", (0, 0), (-1, -1), 9),
                            ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                            ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
                            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ]
                    )
                )
                story.append(table)
            story.append(Spacer(1, 3 * mm))

        if not has_content:
            story.append(Paragraph("本年度暂无相关数据。", body_style))

        buf = BytesIO()
        SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=20 * mm,
            rightMargin=20 * mm,
            topMargin=20 * mm,
            bottomMargin=20 * mm,
            title=title,
        ).build(story)
        return buf.getvalue()


# 向后兼容单例
report_export_service = ReportExportService()
